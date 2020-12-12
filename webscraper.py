
import time

import undetected_chromedriver as uc
import requests
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import html2text
from wordautomate import WordDocument

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.
headers = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_4) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/83.0.4103.97 Safari/537.36"}
def scrape_content(results,d):
    for result in results:
        link = result.find_element_by_tag_name("a")
        href = link.get_attribute("href")
        print("Scraping from :" + href)
        try:
            content=requests.get(href,headers=headers)
        except Exception as e:
            print(e)
            continue


        h = html2text.HTML2Text()
        h.ignore_links=True
        h.ignore_images=True


        d.add_heading("Content scraped from :"+href, WD_PARAGRAPH_ALIGNMENT.CENTER)
        d.add_paragraph(h.handle(content.text), WD_PARAGRAPH_ALIGNMENT.LEFT)
        d.add_page_break()
# Press the green button in the gutter to run the script.
def main():
    br=uc.Chrome()
    print("Enter the name of word file")
    wfname=input()
    print("Enter search keywords")
    search=input()
    br.get('https://www.google.com/search?q='+search)
    time.sleep(2)
    wfpath = r'C:\Users\neha\Documents\AllWordFilesHere' + r'\\' + wfname + ".docx"
    d = WordDocument(wfpath)
    for i in range(0,5):
        results = br.find_elements_by_css_selector('div.g')
        try:
            scrape_content(results,d)
        except Exception as e:
            print(e)
            pass
        br.find_element_by_link_text('Next').click()

    del d



    br.close()
    br.quit()

# See PyCharm help at https://www.jetbrains.com/help/pycharm/