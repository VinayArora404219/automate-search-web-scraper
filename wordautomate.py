from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt,Inches

class WordDocument:
    def __init__(self,path):
        """
        :param path:file path along with name where the file must be stored
        """
        self.path=path
        self.document=Document()

    def __del__(self):
        self.document.save(self.path)

    def header_n_footer_style(self):
        style = self.document.styles.add_style('Footer', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Cambria'
        font.size = Pt(12)


    def setFont(self,font_size,font_name='Cambria',rgbcolor=None):
        style = self.document.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)


    def create_table(self,rows,columns,data):
        """
        :param rows:Number of rows
        :param columns:Number of columns
        :param data:list of lists or tuple of tuples for example ((1,2,3),(4,5,6),(7,8,9))
        :return:nothing
        Example:
        records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam'),
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam'),
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )
    create_table(9,3,records)
        """
        table = self.document.add_table(rows=rows, cols=columns)
        for i in range(rows):
            hdr_cells=table.rows[i].cells
            for j in range(columns):
                hdr_cells[j].text=str(data[i][j])

    def add_heading(self,heading,alignment):
        """
        :param document:docx.Document object
        :param heading:heading to be added.Must be of type str.
        :param alignment:must be an integer constant from docx.enum.text
        :return:nothing
        """
        h=self.document.add_heading(heading)
        h.alignment=alignment

    def add_paragraph(self,para_text,alignment):
        """
           :param document:docx.Document object
           :param heading:heading to be added.Must be of type str.
           :param alignment:must be an integer constant from docx.enum.text
           :return:nothing
        """

        p=self.document.add_paragraph(para_text)
        p.style=self.document.styles['Normal']
        p.alignment=alignment

    def add_picture(self,path,height,width):
        """
        :param path:path of image in OS
        :param width:width of image.must be float value
        :return:
        """
        self.document.add_picture(path, width=Inches(width),height=Inches(height))

    def create_header(self,header_text):
        """
        :param document:docx.Document object
        :param header_text:Text which must be put in header.Must be str.\t must be used for alignment
        :return:
        """
        for section in self.document.sections:
            header = section.header
            for paragraph in header.paragraphs:
                paragraph.text = header_text
                paragraph.style = self.document.styles["Header"]

    def create_footer(self,footer_text):
        """
            :param document:docx.Document object
            :param header_text:Text which must be put in footer.Must be str.\t must be used for alignment
            :return:
        """
        for section in self.document.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                paragraph.style = self.document.styles["Footer"]
                paragraph.text = footer_text


    def add_page_break(self):
        self.document.add_page_break()

# d=WordDocument(r'C:\Users\neha\Documents\AllWordFilesHere\tmpfiles\test.docx')
#
# d.create_header("gead")
# d.create_footer("Vinay Arora\t\t08729802018")
# d.setFont(14)
# d.add_heading("This is the heading",WD_PARAGRAPH_ALIGNMENT.CENTER)
# d.add_paragraph("This is a paragraph",WD_PARAGRAPH_ALIGNMENT.LEFT)
# d.add_picture(r'C:\Users\neha\Pictures\index.jpg',2.5,2.5)
# d.add_page_break()
# d.add_heading("Heading 2",WD_PARAGRAPH_ALIGNMENT.CENTER)

